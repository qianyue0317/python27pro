# -*- encoding:utf-8 -*-

from docx import Document

import sys
import re
import os

import shutil
from convertWord2html import wordsToHtml
from convertWord2html import quitWord
from ziputil import zip

originStr1 = '''**'''
originStr2 = '''##D##'''
originStr3 = '''$$0000000$$'''
originStr4 = '''**'''

styleName = "DefaultParagraph"

defaultencoding = 'utf-8'
if sys.getdefaultencoding() != defaultencoding:
    reload(sys)
    sys.setdefaultencoding(defaultencoding)


def dealFunc(root=r'E:\python\Python27Project\aboutWord', fileName=r'0805000.docx'):
    filePath = os.path.join(root, fileName)
    newDir = os.path.join(root, fileName[0:7])
    if os.path.exists(newDir):
        return
    print('filePath是什么:' + filePath)
    # try:
    myDoc = Document(filePath)
    index = 1
    for pg in myDoc.paragraphs:
        if len(pg.text) == 7 and pg.text.isdigit():  # 是纯数字字符串
            pg.text = ' '  # 赋空串
            continue
        # pg.text = pg.text.replace('\n','')
        styleName = pg.style.name
        theStyle = pg.style
        if pg.text.startswith('答案：') or pg.text.startswith('答案:') or pg.text.startswith('答案：'):
            # print(pg.text)
            answer = re.findall('[A-D]', pg.text)
            # print (answer)
            # if len(answer) == 0: # 防止越界
            if answer.count == 0:  # 防止越界
                answer = ['']
            insertedStr3 = originStr3.replace('0000000', fileName[0:7])
            insertedStr2 = originStr2.replace('D', answer[0])
            pg.insert_paragraph_before(originStr1, styleName)
            pg.insert_paragraph_before(insertedStr2, styleName)
            pg.insert_paragraph_before(insertedStr3, styleName)
            pg.insert_paragraph_before(originStr4, styleName)
        text = pg.text
        # print('行:'+text)
        matchstr = str(index) + '.'  # 题目序号
        matchStr2 = str(index) + '．'
        matchStr3 = str(index) + ' .'
        matchStr4 = str(index) + ' ．'
        matchStr5 = '例题' + str(index)
        # print (matchstr)
        if text.startswith(matchstr):
            # print(text)
            replaceStr(pg, matchstr)
            if not index == 1:
                pg.insert_paragraph_before('++++')
            index += 1
        elif text.startswith(matchStr2):
            # print(text)
            replaceStr(pg, matchstr)
            if not index == 1:
                pg.insert_paragraph_before('++++')
            index += 1
        elif text.startswith(matchStr3):
            # print(text)
            replaceStr(pg, matchstr)
            if not index == 1:
                pg.insert_paragraph_before('++++')
            index += 1
        elif text.startswith(matchStr4):
            # print(text)
            replaceStr(pg, matchstr)
            if not index == 1:
                pg.insert_paragraph_before('++++')
            index += 1
        elif text.startswith(matchStr5):
            # print(text)
            replaceStr(pg, matchstr)
            if not index == 1:
                pg.insert_paragraph_before('++++')
            index += 1

    os.mkdir(newDir)
    newFilePath = os.path.join(newDir, fileName[0:7] + '.docx')
    myDoc.save(newFilePath)
    wordsToHtml(newDir, fileName[0:7] + '.docx')
    # os.remove(newFilePath)
    index = 1
    # except Exception as e:
    #     print(e)
    #     print('转换出问题的文件:'+filePath)


def replaceStr(content, matchStr):
    print(content.runs[0].text)
    # 如果run只包含一个数字
    if len(content.runs[0].text) == 1:
        content.runs[0].text = ''
        content.runs[1].text = content.runs[1].text[1:]
    else:
        try:
            content.runs[0].text = content.runs[0].text[len(matchStr) + 1:]
        except:
            content.runs[0].text = content.runs[0].text[2:]


def isNewTi(title=''):
    if title == None:
        return False
    enPoint = title.find('.')
    if title[:enPoint].isdigit():
        return True
    return False


if '__main__' == __name__:
    dealFunc()
    quitWord()

    '''
    filePath = os.path.join(root, fileName)
    # print('filePath是什么:'+filePath)
    try:
        myDoc = Document(filePath)
        for ps in myDoc.paragraphs:
            styleName = ps.style.name
            if ps.text.startswith('答案：') or ps.text.startswith('答案:'):
                insertedStr = originStr.replace('0000000',fileName[0:7])
                insertedStr = insertedStr.replace('D', ps.text[-1:])
                ps.insert_paragraph_before(insertedStr)
            text = ps.text
            # if re.match('[0-9]*]',text):
            print(text)
            if isNewTi(text):
                ps.insert_paragraph_before('++++\n')
                ps.text = text[2:]
                print('是题目的开头:'+text)
        newDir = os.path.join(root, fileName[0:7])
        os.mkdir(newDir)
        newFilePath = os.path.join(newDir, fileName[0:7]+'.docx')
        myDoc.save(newFilePath)
        wordsToHtml(newDir,fileName[0:7]+'.docx')
        os.remove(newFilePath)
    except BaseException as e:
        print(e)
        print('转换出问题的文件:'+filePath)
    '''
